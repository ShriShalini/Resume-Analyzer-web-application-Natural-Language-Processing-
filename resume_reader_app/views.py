from django.shortcuts import render
from django.http import HttpResponse
from django.core.files.storage import FileSystemStorage
import docx2txt,re,docx,minecart
import spacy,io
import re
import docx
import csv
from docx.opc.constants import RELATIONSHIP_TYPE as RT

import tabula
from tika import parser
from collections import Counter
import en_core_web_sm,PyPDF2,os,comtypes.client,time
from docx2pdf import convert
from docx2python import docx2python
#import minecart
from pprint import pprint
from PyPDF2 import PdfFileReader
#import pdftotext
nlp = en_core_web_sm.load()
import minerpdf
from minerpdf.pdfparser import PDFParser
from minerpdf.pdfdocument import PDFDocument
from minerpdf.pdfpage import PDFPage
from minerpdf.pdfinterp import PDFResourceManager, PDFPageInterpreter
from minerpdf.converter import PDFPageAggregator
from minerpdf.layout import LAParams, LTTextBox, LTTextLine,LTChar


def extract_name(resume,celltext):
    doc  = docx.Document(resume)
    all_paras = doc.paragraphs
    para_text = []
    name_list_regex = []
    i = 1
    name_in_doc_table = ''
    celltext = []
    flag=0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                r = re.compile(r'[A-Za-z.]+[\s]?[A-Za-z]*\s[A-Za-z]*')
                #if(i<2):
                w = r.findall(cell.text)
                if(name_in_doc_table=='' and len(w)==1):
                    if('resume' not in w[0].lower() and 'name' not in w[0].lower()):
                        flag=1
                        name_in_doc_table = w[0]
                        break
            if(flag==1):
                break
        if(flag==1):
            break



    # print('nameindoctable',name_in_doc_table)
    name_in_doc = ''
    for para in all_paras:
        r = re.compile(r'[A-Za-z.]+[\s]?[A-Za-z]*\s[A-Za-z]*')
        w = r.findall(para.text)
        if(name_in_doc=='' and len(w)==1):
            # print('h',w)
            if('resume' not in w[0].lower() and 'name' not in w[0].lower()):
                name_in_doc = w[0]
        para_text.append(para.text)

    # print('nameindoc',name_in_doc)

    # print('name list',name_list_regex)
    para_text_string = ('   ,   '.join(para_text))
    doc = nlp(para_text_string)
    name_spacy = ''
    for X in doc.ents:
        # print(X)
        if(X.label_=='PERSON'):
            name_spacy = X.text
    print('nameindoctable',name_in_doc_table,'nameindoc',name_in_doc,'namespacy',name_spacy)
    if(name_in_doc_table.find('\n')!=-1):
        name_in_doc_table = name_in_doc_table[:name_in_doc_table.find('\n')]
    if(name_in_doc.find('\n')!=-1):
        name_in_doc = name_in_doc[:name_in_doc.find('\n')]
    if(name_spacy.find('\n')!=-1):
        name_spacy = name_spacy[:name_spacy.find('\n')]
    return [name_in_doc_table,name_in_doc,name_spacy]

def extract_email_addresses_docs(resume,celltext):
    r = re.compile(r'\S+@\S+')
    doc  = docx.Document(resume)
    rels = doc.part.rels
    # print("r")
    w_iterrels = []
    mail_link = ''
    for rel in rels:
        if rels[rel].reltype == RT.HYPERLINK:
            link = (rels[rel]._target)
            # print('iteresmail',link)

            if('linkedin' in link.lower()):
                w_iterrels.append(link)
            elif('@' in link.lower()):
                mail_link = link
    all_paras = doc.paragraphs
    para_text = []
    for para in all_paras:
        para_text.append(para.text)
    para_text_string = (' , '.join(para_text))
    w=[]
    w = r.findall(para_text_string)
    w_cell = r.findall(celltext)
    return [mail_link,w_iterrels]

def extract_phone_numbers_doc(resume,celltext):
    doc  = docx.Document(resume)
    all_paras = doc.paragraphs
    para_text = []
    phone_numbers = []
    max_len = 0
    max_cell = ''
    s=r'([\+\(]*[\s\d\.\-\(\)]*)'
    r = re.compile(s)
    phone_numbers_cells = (r.findall(celltext))
    for i in range(0,len(phone_numbers_cells)):
        item = phone_numbers_cells[i].replace(' ','').replace('\t','')
        if(len(item)>max_len):
            max_len = len(phone_numbers_cells[i])
            max_cell  = phone_numbers_cells[i]

    max_len = 0
    max = ''
    for para in all_paras:
        para_text.append(para.text)
        phone_numbers = (r.findall(para.text))

        for i in range(0,len(phone_numbers)):
            item = phone_numbers[i].replace(' ','').replace('\t','')
            if(len(item)>max_len):
                max_len = len(phone_numbers[i])
                max  = phone_numbers[i]


    return [max,max_cell]




def font_type_and_font_size_extraction_doc(resume,celltext):
    doc = docx.Document(resume)
    l = []
    font_dictionary = {}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        fonts = run.font
                        if fonts.size is not None:
                            if(fonts.name not in font_dictionary):
                                if(fonts.name is not None):
                                    set_font = set()
                                    set_font.add(fonts.size.pt)
                                    font_dictionary[fonts.name] = set_font
                            else:
                                if(fonts.name is not None):
                                    font_dictionary[fonts.name].add(fonts.size.pt)

                        else:
                            if(fonts.name not in font_dictionary):
                                if(fonts.name is not None):
                                    set_font = set()
                                    set_font.add('Font size is below 12.0')
                                    font_dictionary[fonts.name] = set_font
                            else:
                                if(fonts.name is not None):
                                    font_dictionary[fonts.name].add('Font size is below 12.0')
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            fonts = run.font

            if fonts.size is not None:
                if(fonts.name not in font_dictionary):
                    if(fonts.name is not None):
                        set_font = set()
                        set_font.add(fonts.size.pt)
                        font_dictionary[fonts.name] = set_font
                else:
                    if(fonts.name is not None):
                        font_dictionary[fonts.name].add(fonts.size.pt)

            else:
                if(fonts.name not in font_dictionary):
                    if(fonts.name is not None):
                        set_font = set()
                        set_font.add('Font size is below 12.0')
                        font_dictionary[fonts.name] = set_font
                else:
                    if(fonts.name is not None):
                        font_dictionary[fonts.name].add('Font size is below 12.0')

    return font_dictionary


def document_num_lines(resume,celltext):
    doc  = docx.Document(resume)
    all_paras = doc.paragraphs
    sum_lines = 0
    for para in all_paras:
        if(re.search('[a-zA-Z]', para.text)):
            sum_lines+=1
    return sum_lines

def extract_table_image_count(resume,celltext):
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    file_path = (os.path.join(BASE_DIR,"media"))
    file_path_for_images = os.path.join(file_path,resume.name)
    doc = docx.Document(resume)
    doc_result = docx2python(file_path_for_images)
    table_count = sum(1 for tab in doc.tables)
    images_count = sum(1 for img in doc_result.images)
    return [table_count,images_count]

"""
-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

"""


def extract_text_characters_count(resume):
    pdfFileObject = open("media\\"+ resume.name, 'rb')

    pdfReader = PyPDF2.PdfFileReader(pdfFileObject)
    email = []
    text=""
    count = []
    page_count = 0
    for page in range(pdfReader.numPages):
        page_count = 0
        pageObject = pdfReader.getPage(page)
        text=pageObject.extractText().replace('\n','').split()
        for t in text:
            page_count+=len(t)

        count.append(page_count)

    return count






def extract_name_pdf(resume):
    pdfFileObject = open("media\\"+ resume.name, 'rb')

    pdfReader = PyPDF2.PdfFileReader(pdfFileObject)
    email = []
    text=""
    single_text = pdfReader.getPage(0).extractText().replace('\n','')
    #print(single_text)
    r = re.compile(r'[A-Za-z.]+[\s]?[A-Za-z]*\s[A-Za-z]*')
    w = r.findall(single_text)
    name = ''
    #print('possibilities',w)
    if(len(w)>1):
        if('resume' not in w[0].lower() and 'name' not in w[0].lower()):
            name = w[0]
        else:
            name = w[1]
    # print('name using regex method',name)

    key = '/Annots'
    uri = '/URI'
    ank = '/A'
    linked_url = ''
    for page in range(pdfReader.numPages):
        pageObject = pdfReader.getPage(page)
        page_url = pageObject.getObject()
        text+=pageObject.extractText().replace('\n','')

        if key in page_url.keys():
            ann = page_url[key]
            for a in ann:
                u = a.getObject()
                if uri in u[ank].keys():
                    if('linkedin' in u[ank][uri]):
                        #print('linkedin' ,u[ank][uri])
                        linked_url = u[ank][uri]
    doc = nlp(text)
    name_list = []
    for X in doc.ents:
        # print()
        #print(X)
        if(X.label_=='PERSON'):
            name_list.append(X.text)
    #print(name,name_list)
    return [name,name_list,linked_url]

def extract_phone_numbers_pdf(resume):
    pdfFileObject = open("media\\"+ resume.name, 'rb')

    pdfReader = PyPDF2.PdfFileReader(pdfFileObject)
    phone_number = []
    text=""
    n_list=[]
    for page in range(pdfReader.numPages):
        pageObject = pdfReader.getPage(page)
        text=pageObject.extractText().replace('\n','')
        #print('text',text)
        s=r'([\+\(]*[\s\d\.\-\(\)]*)'
        r = re.compile(s)

        numbers = r.findall(text)
        # print(numbers)

        for i in numbers:
            i = i.replace(' ','').replace('\t','')
            contains_digit = False
            for character in i:
                if character.isdigit():
                    contains_digit = True
                    break
            if(contains_digit==True):
                n_list.append(i)
        # print('nlist,',n_list)
    pdfFileObject.close()
    if(len(n_list)>0):
        return max(n_list,key=len)
    else:
        return 'Phone number not given'

def extract_font_font_size_pdf(resume):
    fp = open('media\\'+resume.name, 'rb')
    parser = PDFParser(fp)
    doc = PDFDocument(parser)
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    laparams.char_margin = 1.0
    laparams.word_margin = 1.0
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    extracted_text = ''

    fontfont={}
    for page in PDFPage.create_pages(doc):
        interpreter.process_page(page)
        layout = device.get_result()
        for lt_obj in layout:
            if(isinstance(lt_obj,minerpdf.layout.LTTextBoxHorizontal)):
                for ltext in lt_obj:
                    if(isinstance(ltext,minerpdf.layout.LTTextLineHorizontal)):
                        for lchar in ltext:
                            if(isinstance(lchar,minerpdf.layout.LTChar)):
                                if lchar.fontname not in fontfont:
                                    fontfont[lchar.fontname] =[ lchar.fontsize]
                                else:

                                    if lchar.fontsize not in fontfont[lchar.fontname]:
                                        fontfont[lchar.fontname].append(lchar.fontsize)

    return fontfont

def extract_email_addresses_pdf(resume):

    pdfFileObject = open("media\\"+ resume.name, 'rb')

    pdfReader = PyPDF2.PdfFileReader(pdfFileObject)
    email = []
    text=""

    for page in range(pdfReader.numPages):
        pageObject = pdfReader.getPage(page)
        text=pageObject.extractText().replace('\n','')

        email.append(re.findall('\S+@\S+', text))
    # print('email',email)

    pdfFileObject.close()

    return email

def extract_table_image_count_pdf(resume):
    df = tabula.read_pdf('media\\'+resume.name,pages="all",multiple_tables=True)

    # print('number of table',len(df))
    pdffile = open("media\\"+ resume.name, 'rb')
    doc = minecart.Document(pdffile)

    #page = doc.get_page(0) # getting a single page

    #iterating through all pages
    images_count = 0
    for page in doc.iter_pages():
        im = page.images
        images_count+=len(im)

    # print('images count',images_count)
    return [len(df),images_count]

def index(request):

    if(request.method == 'POST' and request.FILES['resume']):

        resume = request.FILES['resume']
        # print('resume',type(resume),resume.name)
        file_Storage = FileSystemStorage()
        file_Name = file_Storage.save(resume.name,resume)
        uploaded_file_url = file_Storage.url(file_Name)
        # print(uploaded_file_url,type(resume),resume.name)
        if('.pdf' in str(resume)):
            # print('It is a pdf')

            email_address_pdf = extract_email_addresses_pdf(resume)
            # print('email_address_from_pdf',email_address_pdf)
            name_pdf = extract_name_pdf(resume)
            # print('primary name possibilities',name_pdf[0])
            # print('secondary name possibilities',name_pdf[1])
            phone_numbers_pdf = extract_phone_numbers_pdf(resume)
            print(phone_numbers_pdf.replace('-',''),'contact no.')
            # print('linkedin url',name_pdf[2])
            #extract_num_lines_pdf(resume)
            text_char_count_pdf  = extract_text_characters_count(resume)
            # print('text_char_count_pdf',text_char_count_pdf)
            font_font_size = extract_font_font_size_pdf(resume)
            # print('font_font_size_pdf',font_font_size)
            table_image_count = extract_table_image_count_pdf(resume)
            # print(table_image_count)
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = 'attachment;filename="resume_details.csv"'

            writer = csv.writer(response)
            writer.writerow(['First row','Primary possibilities for name','Secondary possibilities for name','Email Address','Phone Number','LinkedIn Link','Number of text characters per page','Font details','Number of tables','Number of images'])
            writer.writerow(['Second row',name_pdf[0],name_pdf[1],email_address_pdf,'Phone number: '+phone_numbers_pdf.replace('-',''),name_pdf[2],text_char_count_pdf,font_font_size,table_image_count[0],table_image_count[1]])
            return response
        elif('.docx' in str(resume)):
                # print('It is a doc')
                doc = docx.Document(resume)
                celltext=''
                rowsum  = 0
                for table in doc.tables:
                    for row in table.rows:
                        rowsum+=1
                        for cell in row.cells:
                            celltext+=cell.text
                num_lines_doc = 0
                num_lines_doc = document_num_lines(resume,celltext)
                num_lines_doc+=rowsum
                # print('number of lines in word document',num_lines_doc)
                phone_number_doc= extract_phone_numbers_doc(resume,celltext)
                phone_number_final =''
                if(len(phone_number_doc[0])>0):

                    phone_number_final = 'Phone number'+phone_number_doc[0]
                else:
                    phone_number_final = 'Phone number'+phone_number_doc[1]
                # print('phone_number_doc',phone_number_final.replace('-',''))

                email_address_doc = extract_email_addresses_docs(resume,celltext)
                # print('linkedin',email_address_doc[1])
                email_address_final = email_address_doc[0]

                # print('email_address_docs',email_address_final)
                name_doc = extract_name(resume,celltext)
                name_doc_final = []
                for i in range(len(name_doc)):
                    if(len(name_doc[1])>0 ):
                        name_doc_final.append(name_doc[1])
                    elif(len(name_doc[0])>0 or len(name_doc[2])>0):
                        if(name_doc[0] not in name_doc_final):
                            name_doc_final.append(name_doc[0])
                        if(name_doc[2] not in name_doc_final):
                            name_doc_final.append(name_doc[2])
                final_name = set(name_doc_final)
                # print('name_Doc',final_name)


                font_doc = font_type_and_font_size_extraction_doc(resume,celltext)
                # print(font_doc)
                table_count_doc,image_count_doc  = extract_table_image_count(resume,celltext)
                # print('table count_doc,image count_doc respectively',table_count_doc,image_count_doc)
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment;filename="resume_details.csv"'

                writer = csv.writer(response)
                writer.writerow(['Primary possibilities for name','Secondary possibilities for name','Email Address','Phone Number','LinkedIn Link','Number of text lines','Font details','Number of tables','Number of images'])
                writer.writerow([final_name,final_name,email_address_final,phone_number_final.replace('-',''),email_address_doc[1],num_lines_doc,font_doc,table_count_doc,image_count_doc])

                return response


    return render(request,'resume_reader_app/index.html')
